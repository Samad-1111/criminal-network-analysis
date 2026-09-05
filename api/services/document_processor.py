"""Document content extraction service.

Supports:
    - TXT  — read raw bytes, decode UTF-8 with Latin-1 fallback
    - CSV  — stdlib csv reader; rows joined to structured text
    - PDF  — pypdf page-by-page text extraction
    - DOCX — python-docx paragraph and table cell extraction
"""

import csv
import io
from pathlib import Path


class DocumentProcessingError(Exception):
    """Raised when document content extraction fails."""


def _extract_txt(file_path: Path) -> str:
    """Extract plain text from a TXT file with UTF-8 → Latin-1 fallback."""
    raw = file_path.read_bytes()
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("latin-1", errors="replace")


def _extract_csv(file_path: Path) -> str:
    """Extract rows from a CSV file and join them as tab-separated lines."""
    raw = file_path.read_bytes()
    text = raw.decode("utf-8", errors="replace")
    lines: list[str] = []
    reader = csv.reader(io.StringIO(text))
    for row in reader:
        lines.append("\t".join(row))
    return "\n".join(lines)


def _extract_pdf(file_path: Path) -> str:
    """Extract text from all pages of a PDF using pypdf."""
    try:
        from pypdf import PdfReader  # deferred import — optional dep at runtime
    except ImportError as exc:
        raise DocumentProcessingError(
            "pypdf is not installed. Install it with: pip install pypdf"
        ) from exc

    reader = PdfReader(str(file_path))
    if len(reader.pages) == 0:
        raise DocumentProcessingError("PDF has no pages.")

    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            page_text = f"[Page {i + 1} extraction error: {exc}]"
        parts.append(page_text)

    extracted = "\n".join(parts).strip()
    if not extracted:
        raise DocumentProcessingError(
            "PDF text extraction produced no content. "
            "The file may be image-only or encrypted."
        )
    return extracted


def _extract_docx(file_path: Path) -> str:
    """Extract text from paragraphs and tables in a DOCX file."""
    try:
        from docx import Document  # deferred import — optional dep at runtime
    except ImportError as exc:
        raise DocumentProcessingError(
            "python-docx is not installed. Install it with: pip install python-docx"
        ) from exc

    doc = Document(str(file_path))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells]
            row_text = "\t".join(row_cells)
            if row_text.strip():
                parts.append(row_text)

    extracted = "\n".join(parts).strip()
    if not extracted:
        raise DocumentProcessingError(
            "DOCX extraction produced no content. The document may be empty."
        )
    return extracted


# Dispatch table: file_type (lowercase, no dot) -> extractor function
_EXTRACTORS = {
    "txt": _extract_txt,
    "csv": _extract_csv,
    "pdf": _extract_pdf,
    "docx": _extract_docx,
}


def extract_text(file_path: str, file_type: str) -> str:
    """Extract plain-text content from a stored document file.

    Args:
        file_path: Absolute path to the stored file on disk.
        file_type:  Lowercase file extension without leading dot (e.g. ``"pdf"``).

    Returns:
        Extracted plain-text string (non-empty).

    Raises:
        DocumentProcessingError: If the file is missing, the type is unsupported,
            or extraction fails for any reason.
    """
    path = Path(file_path)
    if not path.exists():
        raise DocumentProcessingError(
            f"File not found at path: {file_path}"
        )
    if not path.is_file():
        raise DocumentProcessingError(
            f"Path is not a regular file: {file_path}"
        )

    normalized_type = file_type.lower().lstrip(".")
    extractor = _EXTRACTORS.get(normalized_type)
    if extractor is None:
        supported = ", ".join(sorted(_EXTRACTORS.keys()))
        raise DocumentProcessingError(
            f"Unsupported file type '{normalized_type}'. "
            f"Supported types: {supported}"
        )

    try:
        return extractor(path)
    except DocumentProcessingError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise DocumentProcessingError(
            f"Unexpected error extracting '{normalized_type}' content: {exc}"
        ) from exc
